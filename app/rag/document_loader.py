from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document as DocxDocument
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


@dataclass
class LoadedDocument:
    source: str
    text: str


def load_documents(data_dir: Path) -> list[LoadedDocument]:
    documents: list[LoadedDocument] = []
    for file_path in sorted(data_dir.rglob("*")):
        if not file_path.is_file() or file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        text = load_single_document(file_path)
        if text:
            documents.append(LoadedDocument(source=str(file_path), text=text))
    return documents


def load_single_document(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".txt":
        return _load_txt_with_fallback(file_path)
    if suffix == ".pdf":
        return _load_pdf(file_path)
    if suffix == ".docx":
        return _load_docx(file_path)
    raise ValueError(f"Unsupported file type: {file_path.suffix}")


def _load_txt_with_fallback(file_path: Path) -> str:
    encodings = ["utf-8", "gbk", "gb18030", "utf-16", "latin-1"]
    for enc in encodings:
        try:
            text = file_path.read_text(encoding=enc)
            return _normalize_text(text)
        except (UnicodeDecodeError, ValueError):
            continue
    return _normalize_text(file_path.read_text(encoding="utf-8", errors="ignore"))


def _load_pdf(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return _normalize_text("\n".join(pages))


def _load_docx(file_path: Path) -> str:
    document = DocxDocument(str(file_path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    return _normalize_text("\n".join(paragraphs))


def _normalize_text(text: str) -> str:
    lines: Iterable[str] = (line.strip() for line in text.splitlines())
    non_empty_lines = [line for line in lines if line]
    return "\n".join(non_empty_lines)
